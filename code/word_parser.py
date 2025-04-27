# -*- coding: utf-8 -*-
import json
from docx import Document
import os
from typing import Dict, Any

class WordParser:
    """用于解析 Word 文档并将其转换为 JSON 格式的类，仅提取文本内容。"""

    def __init__(self, file_path: str, doc_type: str = "word"):
        """
        初始化 WordParser 类。
        :param file_path: Word 文件路径
        :param doc_type: 文档类型
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Word file not found: {file_path}")

        self.file_path = file_path
        self.file_name = os.path.basename(file_path)
        self.doc_type = doc_type
        self.result = {
            "doc_type": self.doc_type,
            "file_name": self.file_name,
            "content": ""
        }

    def parse(self) -> Dict[str, Any]:
        """
        解析 Word 文档，提取段落和表格的文本内容。
        :return: 解析后的 JSON 数据，包含 doc_type、file_name 和 content
        """
        try:
            # 加载 Word 文档
            doc = Document(self.file_path)

            # 存储所有文本内容的列表
            text_parts = []

            # 提取段落文本
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # 忽略空段落
                    text_parts.append(text)

            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:  # 忽略空行
                        # 用制表符 \t 分隔单元格内容，模拟表格结构
                        text_parts.append("\t".join(row_text))

            # 将所有文本拼接为一个长字符串，用换行符分隔
            self.result["content"] = "\n".join(text_parts)

        except Exception as e:
            raise Exception(f"Error processing Word file: {str(e)}")

        return self.result

    def save_to_file(self, output_dir: str = "llm_output_test"):
        """
        将解析结果保存为 JSON 文件，文件名为“word名称_1.json”。
        :param output_dir: 输出目录
        """
        # 确保输出目录存在
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # 构造输出文件名：word名称_1.json
        base_name = os.path.splitext(self.file_name)[0]  # 去掉扩展名
        output_file_name = f"{base_name}_1.json"
        output_path = os.path.join(output_dir, output_file_name)

        # 保存 JSON 文件
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(self.result, f, ensure_ascii=False, indent=2)
        print(f"Saved Word content to: {output_path}")

def main():
    # 替换为你的 Word 文件路径
    file_path = r"C:\Users\dreame\Desktop\电子元件RAG\文档\扫地机重新烧号软件使用说明.docx"

    # 使用 WordParser 解析 Word 文件
    parser = WordParser(file_path)
    result = parser.parse()

    # 打印结果
    print("=== Parsed Word ===")
    print(f"Doc Type: {result['doc_type']}")
    print(f"File Name: {result['file_name']}")
    print(f"Content:\n{result['content']}")

    # 保存为 JSON 文件
    parser.save_to_file(output_dir="llm_output_test")

if __name__ == "__main__":
    main()